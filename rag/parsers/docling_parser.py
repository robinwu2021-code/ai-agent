"""
rag/parsers/docling_parser.py — Docling 文档解析器（推荐默认）

Docling（IBM 2024 开源）特点：
  - PDF 复杂表格 → 结构化 Markdown（保留行列关系）
  - 扫描件 OCR（EasyOCR / Tesseract）
  - DOCX / XLSX / PPTX / HTML / CSV 全格式支持
  - 输出统一 Markdown，与框架完全无关

安装：pip install docling
降级回退：python-docx / pypdf / openpyxl（仅需其中相关包）
"""
from __future__ import annotations

import asyncio
import csv
import io
from pathlib import Path

import structlog

from rag.parsers.base import ParsedDocument

log = structlog.get_logger(__name__)

_SUPPORTED = frozenset({
    ".pdf", ".docx", ".xlsx", ".xls", ".pptx",
    ".html", ".htm", ".md", ".markdown", ".csv", ".txt",
})


class DoclingParser:
    """基于 Docling 的多格式文档解析器。"""

    def __init__(
        self,
        do_ocr: bool = True,
        do_table_structure: bool = True,
        table_mode: str = "accurate",
        ocr_lang: list[str] | None = None,
    ) -> None:
        self._do_ocr = do_ocr
        self._do_table_structure = do_table_structure
        self._table_mode = table_mode
        self._ocr_lang = ocr_lang or ["zh", "en"]
        self._converter = None

    def _get_converter(self):
        if self._converter is not None:
            return self._converter
        try:
            from docling.document_converter import DocumentConverter, PdfFormatOption
            from docling.datamodel.pipeline_options import PdfPipelineOptions
            opts = PdfPipelineOptions(
                do_ocr=self._do_ocr,
                do_table_structure=self._do_table_structure,
            )
            self._converter = DocumentConverter(
                format_options={"pdf": PdfFormatOption(pipeline_options=opts)}
            )
            log.info("docling.converter_ready")
        except ImportError:
            log.warning("docling.not_installed", hint="pip install docling")
            self._converter = None
        except Exception as exc:
            log.warning("docling.init_failed", error=str(exc))
            self._converter = None
        return self._converter

    def supports(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in _SUPPORTED

    async def parse(self, file_path: str) -> ParsedDocument:
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix in {".txt", ".md", ".markdown"}:
            return await self._parse_plain(path)
        if suffix == ".csv":
            return await self._parse_csv(path)

        converter = self._get_converter()
        if converter is None:
            return await self._parse_fallback(path)

        try:
            return await asyncio.to_thread(self._run_docling, converter, path)
        except Exception as exc:
            log.warning("docling.convert_failed", path=str(path), error=str(exc))
            return await self._parse_fallback(path)

    # ── Docling 同步转换（在线程池中运行）───────────────────────────

    def _run_docling(self, converter, path: Path) -> ParsedDocument:
        result = converter.convert(str(path))
        doc = result.document
        content = doc.export_to_markdown()

        headings: list[dict] = []
        has_headings = False
        try:
            for item in doc.texts:
                lbl = str(getattr(item, "label", "")).lower()
                if any(x in lbl for x in ("heading", "title", "h1", "h2", "h3")):
                    level = 1
                    if "h2" in lbl or "heading2" in lbl or "section" in lbl:
                        level = 2
                    elif "h3" in lbl or "heading3" in lbl:
                        level = 3
                    text = getattr(item, "text", "").strip()
                    if text:
                        headings.append({"level": level, "text": text})
                        has_headings = True
        except Exception:
            pass

        tables: list[dict] = []
        try:
            for tbl in doc.tables:
                md = (tbl.export_to_markdown()
                      if hasattr(tbl, "export_to_markdown") else str(tbl))
                tables.append({"markdown": md})
        except Exception:
            pass

        page_count = 0
        try:
            page_count = len(doc.pages) if hasattr(doc, "pages") else 0
        except Exception:
            pass

        if not has_headings:
            has_headings = "# " in content or "## " in content

        return ParsedDocument(
            content=content,
            source=str(path),
            format=path.suffix.lower().lstrip("."),
            has_headings=has_headings,
            title=headings[0]["text"] if headings else path.stem,
            page_count=page_count,
            headings=headings,
            tables=tables,
            metadata={"parser": "docling"},
        )

    # ── 简单格式直读 ──────────────────────────────────────────────

    async def _parse_plain(self, path: Path) -> ParsedDocument:
        try:
            content = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            content = path.read_text(encoding="gbk", errors="replace")

        headings = []
        for line in content.splitlines():
            s = line.lstrip()
            for lvl, prefix in [(1, "# "), (2, "## "), (3, "### ")]:
                if s.startswith(prefix):
                    headings.append({"level": lvl, "text": s[len(prefix):].strip()})
                    break

        return ParsedDocument(
            content=content,
            source=str(path),
            format=path.suffix.lower().lstrip("."),
            has_headings=bool(headings),
            title=headings[0]["text"] if headings else path.stem,
            headings=headings,
            metadata={"parser": "plain"},
        )

    async def _parse_csv(self, path: Path) -> ParsedDocument:
        try:
            raw = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            raw = path.read_text(encoding="gbk", errors="replace")

        reader = csv.reader(io.StringIO(raw))
        rows = list(reader)
        if not rows:
            return ParsedDocument(content="", source=str(path), format="csv")

        header = rows[0]
        lines = [
            "| " + " | ".join(str(c) for c in header) + " |",
            "| " + " | ".join(["---"] * len(header)) + " |",
        ]
        for row in rows[1:]:
            padded = row + [""] * max(0, len(header) - len(row))
            lines.append("| " + " | ".join(str(c) for c in padded[:len(header)]) + " |")

        content = "\n".join(lines)
        return ParsedDocument(
            content=content,
            source=str(path),
            format="csv",
            title=path.stem,
            tables=[{"markdown": content}],
            metadata={"parser": "csv", "rows": len(rows) - 1, "cols": len(header)},
        )

    # ── Docling 不可用时的回退解析 ─────────────────────────────────

    async def _parse_fallback(self, path: Path) -> ParsedDocument:
        log.warning("docling.fallback_parser", path=str(path))
        suffix = path.suffix.lower()
        content = ""

        if suffix == ".docx":
            try:
                import docx as _docx
                doc = await asyncio.to_thread(_docx.Document, str(path))
                parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        style = (para.style.name or "").lower()
                        if "heading 1" in style:
                            parts.append(f"# {para.text}")
                        elif "heading 2" in style:
                            parts.append(f"## {para.text}")
                        elif "heading 3" in style:
                            parts.append(f"### {para.text}")
                        else:
                            parts.append(para.text)
                content = "\n\n".join(parts)
            except Exception as e:
                log.warning("fallback.docx_failed", error=str(e))

        elif suffix == ".pdf":
            try:
                import pypdf
                def _read():
                    r = pypdf.PdfReader(str(path))
                    return "\n\n".join(p.extract_text() or "" for p in r.pages)
                content = await asyncio.to_thread(_read)
            except Exception as e:
                log.warning("fallback.pdf_failed", error=str(e))

        elif suffix in {".xlsx", ".xls"}:
            try:
                import openpyxl
                def _read():
                    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
                    parts = []
                    for sheet in wb.worksheets:
                        parts.append(f"## {sheet.title}")
                        rows = list(sheet.iter_rows(values_only=True))
                        if rows:
                            hdr = [str(c) if c is not None else "" for c in rows[0]]
                            parts.append("| " + " | ".join(hdr) + " |")
                            parts.append("| " + " | ".join(["---"] * len(hdr)) + " |")
                            for row in rows[1:]:
                                cells = [str(c) if c is not None else "" for c in row]
                                parts.append("| " + " | ".join(cells[:len(hdr)]) + " |")
                    return "\n".join(parts)
                content = await asyncio.to_thread(_read)
            except Exception as e:
                log.warning("fallback.xlsx_failed", error=str(e))

        has_headings = "## " in content or "# " in content
        return ParsedDocument(
            content=content,
            source=str(path),
            format=suffix.lstrip("."),
            has_headings=has_headings,
            title=path.stem,
            metadata={"parser": "fallback"},
        )
