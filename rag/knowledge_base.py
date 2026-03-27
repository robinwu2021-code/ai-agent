"""
rag/knowledge_base.py — RAG 增强检索知识库

功能：
  - DocumentIngester   文档摄入（PDF / TXT / MD / HTML）
  - TextChunker        智能分块（固定窗口 + 语义边界）
  - HybridRetriever    混合检索（向量 + BM25 关键词）
  - Reranker           重排序（交叉编码器 / LLM 打分）
  - KnowledgeBase      统一入口
"""
from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import structlog

log = structlog.get_logger(__name__)


# ─────────────────────────────────────────────
# Data Types
# ─────────────────────────────────────────────

@dataclass
class Document:
    id:       str
    source:   str           # 文件路径或 URL
    content:  str
    metadata: dict = field(default_factory=dict)

@dataclass
class Chunk:
    id:          str
    doc_id:      str
    text:        str
    chunk_index: int
    metadata:    dict = field(default_factory=dict)
    embedding:   list[float] | None = None
    score:       float = 0.0


# ─────────────────────────────────────────────
# Document Ingester
# ─────────────────────────────────────────────

class DocumentIngester:
    """
    从多种格式读取文档内容，统一转为纯文本。
    支持：.txt .md .html .pdf（需 pypdf）.docx（需 python-docx）
    """

    def ingest_file(self, path: str | Path) -> Document:
        p = Path(path)
        if not p.exists():
            raise FileNotFoundError(str(p))
        suffix = p.suffix.lower()
        if suffix in (".txt", ".md"):
            content = p.read_text(errors="replace")
        elif suffix == ".html":
            content = self._strip_html(p.read_text(errors="replace"))
        elif suffix == ".pdf":
            content = self._read_pdf(p)
        elif suffix == ".docx":
            content = self._read_docx(p)
        else:
            content = p.read_text(errors="replace")
        doc_id = hashlib.md5(str(p).encode()).hexdigest()[:12]
        log.info("rag.ingest", source=str(p), chars=len(content))
        return Document(id=doc_id, source=str(p), content=content,
                        metadata={"filename": p.name, "suffix": suffix})

    def ingest_text(self, text: str, source: str = "inline") -> Document:
        doc_id = hashlib.md5(text[:200].encode()).hexdigest()[:12]
        return Document(id=doc_id, source=source, content=text)

    def _strip_html(self, html: str) -> str:
        text = re.sub(r"<style[^>]*>.*?</style>", "", html, flags=re.DOTALL)
        text = re.sub(r"<script[^>]*>.*?</script>", "", text, flags=re.DOTALL)
        text = re.sub(r"<[^>]+>", " ", text)
        text = re.sub(r"\s+", " ", text)
        return text.strip()

    def _read_pdf(self, path: Path) -> str:
        """
        提取 PDF 文本。

        策略（按优先级）：
        1. pdfplumber — 逐页提取，每页之间用 \\n\\n 分隔，保留段落边界
        2. pypdf      — 回退，同样按页用 \\n\\n 拼接（避免全文单 \\n 导致无法分块）
        """
        # ── 优先用 pdfplumber（段落结构更完整）──────────────────────────
        try:
            import pdfplumber
            pages: list[str] = []
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    text = text.strip()
                    if text:
                        pages.append(text)
            if pages:
                return "\n\n".join(pages)
        except Exception:
            pass  # 回退到 pypdf

        # ── 回退：pypdf（页间用 \n\n，避免 \n 拼接导致零分块）──────────
        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            pages = []
            for page in reader.pages:
                text = (page.extract_text() or "").strip()
                if text:
                    pages.append(text)
            return "\n\n".join(pages)
        except ImportError:
            raise RuntimeError("请安装 pypdf：pip install pypdf")

    def _read_docx(self, path: Path) -> str:
        try:
            import docx
            doc = docx.Document(str(path))
            return "\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            raise RuntimeError("请安装 python-docx：pip install python-docx")


# ─────────────────────────────────────────────
# Text Chunker
# ─────────────────────────────────────────────

class TextChunker:
    """
    将长文档分割为适合检索的块。
    策略：按段落边界切割，块大小在 [min_size, max_size] 之间，
    相邻块有 overlap 字符重叠以保持上下文连贯。
    """

    def __init__(
        self,
        max_chars: int = 800,
        min_chars: int = 200,
        overlap:   int = 100,
    ) -> None:
        self._max     = max_chars
        self._min     = min_chars
        self._overlap = overlap

    def chunk(self, doc: Document) -> list[Chunk]:
        paragraphs = self._split_paragraphs(doc.content)
        chunks: list[Chunk] = []
        buf = ""
        idx = 0

        for para in paragraphs:
            if len(buf) + len(para) <= self._max:
                buf += para + "\n"
            else:
                if len(buf) >= self._min:
                    chunks.append(self._make_chunk(doc, buf.strip(), idx))
                    idx += 1
                    buf = buf[-self._overlap:] + para + "\n"
                else:
                    buf += para + "\n"

        if buf.strip():
            chunks.append(self._make_chunk(doc, buf.strip(), idx))

        log.debug("rag.chunked", doc_id=doc.id, chunks=len(chunks))
        return chunks

    def _split_paragraphs(self, text: str) -> list[str]:
        """
        段落识别策略（三级降级）：

        1. 双换行 / 句号+换行 → 段落边界（中英文通用）
        2. 超出 max_chars 的块：单换行进一步切分（处理英文 PDF 无空行情况）
        3. 超长单行：按句子边界强制切断
        """
        # 第一轮：双换行 + 中英文句尾+换行
        raw = re.split(r"\n{2,}|(?<=[。！？.!?])\n", text)

        result: list[str] = []
        for block in raw:
            block = block.strip()
            if not block:
                continue

            if len(block) <= self._max:
                result.append(block)
            else:
                # 第二轮：英文 PDF 段落内只有单换行 → 按 \n 切分
                lines = [ln.strip() for ln in block.split("\n") if ln.strip()]
                buf = ""
                for line in lines:
                    if len(buf) + len(line) + 1 <= self._max:
                        buf = (buf + " " + line).strip() if buf else line
                    else:
                        if buf:
                            result.append(buf)
                        # 第三轮：单行本身超长 → 按句子边界切断
                        if len(line) > self._max:
                            result.extend(self._split_long_line(line))
                            buf = ""
                        else:
                            buf = line
                if buf:
                    result.append(buf)

        return result

    def _split_long_line(self, line: str) -> list[str]:
        """将单行超长文本按句子边界切分为不超过 max_chars 的片段。"""
        parts = re.split(r"(?<=[.!?])\s+", line)
        chunks: list[str] = []
        buf = ""
        for part in parts:
            candidate = (buf + " " + part).strip() if buf else part
            if len(candidate) <= self._max:
                buf = candidate
            else:
                if buf:
                    chunks.append(buf)
                buf = part if len(part) <= self._max else part[: self._max]
        if buf:
            chunks.append(buf)
        return chunks or [line[: self._max]]

    def _make_chunk(self, doc: Document, text: str, idx: int) -> Chunk:
        cid = f"{doc.id}_c{idx:04d}"
        return Chunk(
            id=cid, doc_id=doc.id, text=text, chunk_index=idx,
            metadata={**doc.metadata, "source": doc.source},
        )


# ─────────────────────────────────────────────
# BM25 Index（轻量实现，无外部依赖）
# ─────────────────────────────────────────────

class BM25Index:
    """简单 BM25 关键词检索，中英文均支持（字符级 n-gram 分词）。"""

    def __init__(self, k1: float = 1.5, b: float = 0.75) -> None:
        self._k1 = k1
        self._b  = b
        self._chunks: list[Chunk] = []
        self._tf:   list[dict[str, float]] = []
        self._df:   dict[str, int]         = {}
        self._avgdl = 0.0

    def build(self, chunks: list[Chunk]) -> None:
        self._chunks = chunks
        self._tf     = []
        self._df     = {}
        total_len = 0
        for chunk in chunks:
            tokens = self._tokenize(chunk.text)
            total_len += len(tokens)
            tf: dict[str, float] = {}
            for t in tokens:
                tf[t] = tf.get(t, 0) + 1
            self._tf.append(tf)
            for t in set(tokens):
                self._df[t] = self._df.get(t, 0) + 1
        self._avgdl = total_len / max(len(chunks), 1)

    def search(self, query: str, top_k: int = 10) -> list[tuple[float, Chunk]]:
        import math
        n = len(self._chunks)
        q_tokens = self._tokenize(query)
        scores   = []
        for i, chunk in enumerate(self._chunks):
            dl  = sum(self._tf[i].values())
            sc  = 0.0
            for t in q_tokens:
                if t not in self._tf[i]:
                    continue
                tf  = self._tf[i][t]
                df  = self._df.get(t, 0)
                idf = math.log((n - df + 0.5) / (df + 0.5) + 1)
                sc += idf * (tf * (self._k1 + 1)) / (
                    tf + self._k1 * (1 - self._b + self._b * dl / self._avgdl)
                )
            scores.append((sc, chunk))
        scores.sort(key=lambda x: x[0], reverse=True)
        return [s for s in scores[:top_k] if s[0] > 0]

    @staticmethod
    def _tokenize(text: str) -> list[str]:
        # 字符 bigram + 单词切分（兼容中英文）
        text  = text.lower()
        words = re.findall(r"[a-z0-9]+", text)
        cjk   = re.findall(r"[\u4e00-\u9fff]", text)
        bigrams = [text[i:i+2] for i in range(len(text)-1)
                   if text[i:i+2].strip() and not text[i:i+2].isspace()]
        return words + cjk + bigrams[:50]  # 限制 bigram 数量


# ─────────────────────────────────────────────
# Hybrid Retriever
# ─────────────────────────────────────────────

class HybridRetriever:
    """
    混合检索：向量检索（语义） + BM25（关键词），RRF 融合排序。
    embed_fn 传入 LLMEngine.embed 方法。
    """

    def __init__(self, embed_fn: Any, rrf_k: int = 60) -> None:
        self._embed   = embed_fn
        self._rrf_k   = rrf_k
        self._chunks: list[Chunk] = []
        self._embeddings: list[list[float]] = []
        self._bm25 = BM25Index()

    async def add_chunks(self, chunks: list[Chunk]) -> None:
        for chunk in chunks:
            try:
                vec = await self._embed(chunk.text)
                chunk.embedding = vec
                self._embeddings.append(vec)
            except Exception as exc:
                # embed 失败（如 Ollama 无 embedding 模型）时跳过向量，仍走 BM25
                log.warning("rag.add_chunk.embed_failed_bm25_only",
                            chunk_id=chunk.id, error=str(exc))
                chunk.embedding = None
            self._chunks.append(chunk)
        self._bm25.build(self._chunks)
        log.info("rag.indexed", total_chunks=len(self._chunks))

    async def search(self, query: str, top_k: int = 5) -> list[Chunk]:
        if not self._chunks:
            return []

        # 向量检索（embed 失败时降级 BM25-only，不抛异常）
        vec_hits: list[tuple[float, "Chunk"]] = []
        try:
            q_vec    = await self._embed(query)
            vec_hits = self._vector_search(q_vec, top_k * 2)
        except Exception as exc:
            log.warning("rag.search.embed_failed_bm25_only", error=str(exc))

        bm25_hits = self._bm25.search(query, top_k * 2)

        # RRF 融合
        scores: dict[str, float] = {}
        chunk_map: dict[str, "Chunk"] = {}
        for rank, (_, chunk) in enumerate(vec_hits):
            scores[chunk.id]    = scores.get(chunk.id, 0) + 1 / (self._rrf_k + rank + 1)
            chunk_map[chunk.id] = chunk
        for rank, (_, chunk) in enumerate(bm25_hits):
            scores[chunk.id]    = scores.get(chunk.id, 0) + 1 / (self._rrf_k + rank + 1)
            chunk_map[chunk.id] = chunk

        ranked = sorted(scores.items(), key=lambda x: x[1], reverse=True)
        results = []
        for cid, sc in ranked[:top_k]:
            c = chunk_map[cid]
            c.score = sc
            results.append(c)
        return results

    def _vector_search(
        self, query_vec: list[float], top_k: int
    ) -> list[tuple[float, Chunk]]:
        import math
        def cosine(a: list[float], b: list[float]) -> float:
            dot  = sum(x * y for x, y in zip(a, b))
            na   = math.sqrt(sum(x*x for x in a))
            nb   = math.sqrt(sum(x*x for x in b))
            return dot / (na * nb) if na * nb > 0 else 0.0
        scored = [(cosine(query_vec, c.embedding or []), c)
                  for c in self._chunks if c.embedding]
        scored.sort(key=lambda x: x[0], reverse=True)
        return scored[:top_k]


# ─────────────────────────────────────────────
# Reranker
# ─────────────────────────────────────────────

class LLMReranker:
    """
    用 LLM 对初步召回的 Chunk 进行相关性重排序。
    生产环境可替换为 Cohere Rerank API 或 Cross-Encoder 模型。
    """

    RERANK_PROMPT = """给出以下查询和文档片段的相关性分数（0-10），只返回数字。
查询：{query}
文档：{text}
分数："""

    def __init__(self, llm_engine: Any) -> None:
        self._llm = llm_engine

    async def rerank(
        self, query: str, chunks: list[Chunk], top_k: int = 3
    ) -> list[Chunk]:
        if not chunks:
            return []
        scored = []
        for chunk in chunks:
            prompt = self.RERANK_PROMPT.format(
                query=query, text=chunk.text[:300]
            )
            raw = await self._llm.summarize(prompt, max_tokens=5)
            try:
                score = float(re.search(r"[\d.]+", raw).group())
            except Exception:
                score = chunk.score * 10
            scored.append((score, chunk))
        scored.sort(key=lambda x: x[0], reverse=True)
        result = [c for _, c in scored[:top_k]]
        for i, c in enumerate(result):
            c.score = scored[i][0] / 10.0
        return result


# ─────────────────────────────────────────────
# Knowledge Base (façade)
# ─────────────────────────────────────────────

class KnowledgeBase:
    """
    RAG 知识库统一入口。

    用法：
        kb = KnowledgeBase(embed_fn=llm.embed, llm_engine=llm)
        await kb.add_file("docs/manual.pdf")
        chunks = await kb.query("如何重置密码", top_k=3)
        context = kb.format_context(chunks)
    """

    def __init__(
        self,
        embed_fn: Any,
        llm_engine: Any | None = None,
        use_reranker: bool = False,
        chunk_max_chars: int = 800,
    ) -> None:
        self._ingester  = DocumentIngester()
        self._chunker   = TextChunker(max_chars=chunk_max_chars)
        self._retriever = HybridRetriever(embed_fn=embed_fn)
        self._reranker  = LLMReranker(llm_engine) if (use_reranker and llm_engine) else None
        self._doc_count = 0

    async def add_file(self, path: str | Path) -> int:
        doc    = self._ingester.ingest_file(path)
        chunks = self._chunker.chunk(doc)
        await self._retriever.add_chunks(chunks)
        self._doc_count += 1
        return len(chunks)

    async def add_text(self, text: str, source: str = "inline") -> int:
        doc    = self._ingester.ingest_text(text, source)
        chunks = self._chunker.chunk(doc)
        await self._retriever.add_chunks(chunks)
        self._doc_count += 1
        return len(chunks)

    async def query(self, query: str, top_k: int = 5) -> list[Chunk]:
        chunks = await self._retriever.search(query, top_k=top_k * 2)
        if self._reranker:
            chunks = await self._reranker.rerank(query, chunks, top_k=top_k)
        return chunks[:top_k]

    @staticmethod
    def format_context(chunks: list[Chunk]) -> str:
        """将检索结果格式化为可注入 Prompt 的字符串。"""
        if not chunks:
            return ""
        parts = ["## 相关知识库内容\n"]
        for i, chunk in enumerate(chunks, 1):
            source = chunk.metadata.get("filename", chunk.doc_id)
            parts.append(f"[{i}] 来源: {source}\n{chunk.text}\n")
        return "\n".join(parts)

    @property
    def doc_count(self) -> int:
        return self._doc_count

    @property
    def chunk_count(self) -> int:
        return len(self._retriever._chunks)
